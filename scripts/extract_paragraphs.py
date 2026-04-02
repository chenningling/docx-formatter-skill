#!/usr/bin/env python3
"""
extract_paragraphs.py — 提取文档所有段落，输出供 Claude 进行语义分析的列表。

用法：
    python extract_paragraphs.py --input 文档.docx

输出格式（标准输出）：
    [00] <空>
    [01] AI的未来发展：机遇与挑战
    [02] <空>
    [03] 人工智能（AI）正以前所未有的速度...（共104字）
    [04] <空>
    [05] 一、技术演进趋势
    ...

Claude 读取此输出后，应判断每个非空段落的类型，并生成 heading_map JSON，
传给 format_docx.py 的 --heading-map 参数。
"""

import argparse
import sys
from pathlib import Path
from docx import Document


# 段落文字超过此长度时截断显示，并标注总字数
PREVIEW_MAX = 50


def extract(input_path: str):
    doc = Document(input_path)
    total = len(doc.paragraphs)

    print(f"文档共 {total} 个段落（含空行）。以下为完整段落列表：\n")
    print("格式说明：[索引] 段落文字（若超长则截断并标注总字数）")
    print("─" * 60)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if not text:
            print(f"[{i:02d}] <空>")
            continue

        # 超长段落截断显示
        char_count = len(text)
        if char_count > PREVIEW_MAX:
            display = text[:PREVIEW_MAX] + f"...（共{char_count}字）"
        else:
            display = text

        # 附上样式名，帮助 Claude 判断（仅当样式名有参考价值时）
        style_hint = ""
        if style and style.lower() not in ("normal", "a", "默认段落字体", "default paragraph font"):
            style_hint = f"  [样式:{style}]"

        print(f"[{i:02d}] {display}{style_hint}")

    print("─" * 60)
    print(f"\n共 {total} 段。")
    print("\n── Claude 分析指引 ──")
    print("请根据上方段落内容，判断每个段落的层级类型，生成如下 JSON：")
    print('{"索引": "类型", ...}')
    print("类型取值：title（文档总标题）/ h1（章/一级）/ h2（节/二级）/ h3（小节/三级）/ body（正文，默认不需填）")
    print("\n判断参考：")
    print("  - title : 文档最顶层标题，通常是第一个有意义的短段落")
    print("  - h1    : 章节标题，如[第一章 绪论]、[一、总体要求]、[1 Introduction]")
    print("  - h2    : 节标题，如[1.1 研究背景]、[(一)强化组织领导]")
    print("  - h3    : 小节标题，如[1.1.1 深度学习的优势]、[1. 制度制定要求]")
    print("  - body  : 正文段落（字数多、无编号结构的叙述性内容）")
    print("\n只需列出 title/h1/h2/h3 的段落，body 不需要填入 JSON（脚本默认处理）。")
    print("\n示例输出：")
    print('  {"0": "title", "5": "h1", "11": "h1", "14": "h2", "18": "h3"}')


def main():
    parser = argparse.ArgumentParser(
        description="提取文档段落列表，供 Claude 语义分析后生成 heading_map"
    )
    parser.add_argument("--input", required=True, help="输入 .docx 文件路径")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"错误：文件不存在：{args.input}", file=sys.stderr)
        sys.exit(1)

    extract(args.input)


if __name__ == "__main__":
    main()
