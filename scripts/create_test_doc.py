#!/usr/bin/env python3
"""
create_test_doc.py — 生成用于测试的 Word 文档。

用法：
    python create_test_doc.py --type academic --output test_academic.docx
    python create_test_doc.py --type gongwen  --output test_gongwen.docx

生成的文档会刻意使用"杂乱"的格式（随意字体、字号不统一），
以便测试 format_docx.py 是否能正确统一格式化。
"""

import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_academic_test_doc(output_path: str):
    """创建一篇带有混乱格式的学术论文草稿。"""
    doc = Document()

    # 故意混用各种格式，模拟真实用户文档的"脏"状态

    # 文档标题（使用 Title 样式）
    title_para = doc.add_paragraph('基于深度学习的中文文本分类研究', style='Title')

    # 正文摘要（普通段落，随意设了 Arial 字体）
    p = doc.add_paragraph('摘要：本文提出了一种基于 BERT 预训练模型的中文文本分类方法。'
                           '实验结果表明，该方法在多个基准数据集上取得了优秀的性能。')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    doc.add_paragraph('')  # 空行

    # 一级标题
    doc.add_heading('1 绪论', level=1)

    # 正文（混用了 Calibri 和不同字号）
    p = doc.add_paragraph('随着自然语言处理技术的快速发展，深度学习方法在文本分类领域取得了显著进展。'
                           '本研究旨在探索预训练语言模型在中文场景下的应用效果。')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

    # 二级标题
    doc.add_heading('1.1 研究背景', level=2)

    p = doc.add_paragraph('文本分类是自然语言处理中的基础任务之一，在情感分析、垃圾邮件检测、'
                           '新闻分类等领域有着广泛的应用。传统方法依赖人工特征工程，效率较低。')
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 三级标题
    doc.add_heading('1.1.1 深度学习的优势', level=3)

    p = doc.add_paragraph('深度学习模型能够自动提取特征，减少对领域知识的依赖。'
                           'BERT、GPT 等预训练模型的出现进一步推动了该领域的发展。')

    # 另一个一级标题
    doc.add_heading('2 相关工作', level=1)

    p = doc.add_paragraph('早期的文本分类方法主要包括朴素贝叶斯、支持向量机等传统机器学习方法。'
                           'TextCNN [Kim, 2014] 首次将卷积神经网络引入文本分类任务，取得了良好效果。')

    doc.add_heading('2.1 预训练语言模型', level=2)

    p = doc.add_paragraph('BERT（Bidirectional Encoder Representations from Transformers）由 Google 于 2018 年提出，'
                           '通过双向 Transformer 结构在大规模语料上进行预训练，在众多 NLP 下游任务上实现了突破性进展。')
    for run in p.runs:
        run.font.size = Pt(14)  # 故意用了错误的字号

    doc.add_heading('3 方法', level=1)

    p = doc.add_paragraph('本文提出的方法基于 BERT 模型进行微调，在输出层添加全连接层进行分类。'
                           '模型训练采用 AdamW 优化器，学习率设置为 2e-5，批大小为 32。')

    doc.add_heading('3.1 模型架构', level=2)
    p = doc.add_paragraph('模型整体架构如下：输入层接收分词后的文本序列，通过 BERT 编码层获得上下文表示，'
                           '最后通过分类头输出预测结果。')

    doc.add_heading('4 实验与结果', level=1)
    p = doc.add_paragraph('在 THUCNews 数据集上，本文方法达到了 97.2% 的准确率，'
                           '相比基线方法提升了 2.3 个百分点。')

    doc.add_heading('5 结论', level=1)
    p = doc.add_paragraph('本文提出了基于 BERT 的中文文本分类方法，实验验证了其有效性。'
                           '未来工作将探索更轻量级的模型结构以降低计算成本。')

    p = doc.add_paragraph('参考文献')
    for run in p.runs:
        run.font.bold = True

    p = doc.add_paragraph('[1] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of deep bidirectional '
                           'transformers for language understanding[J]. arXiv:1810.04805, 2018.')
    p = doc.add_paragraph('[2] Kim Y. Convolutional neural networks for sentence classification[J]. '
                           'arXiv:1408.5882, 2014.')

    doc.save(output_path)
    print(f'✓ 学术论文测试文档已创建：{output_path}')


def create_gongwen_test_doc(output_path: str):
    """创建一份带有混乱格式的公文草稿。"""
    doc = Document()

    # 公文标题
    title_para = doc.add_paragraph('关于加强网络安全工作的通知', style='Title')

    p = doc.add_paragraph('各部门、各单位：')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    p = doc.add_paragraph('为深入贯彻落实网络安全法律法规，切实加强我单位网络安全工作，'
                           '有效防范网络安全风险，现将有关事项通知如下。')

    # 一级标题
    doc.add_heading('一、总体要求', level=1)

    p = doc.add_paragraph('各部门要高度重视网络安全工作，将其纳入重要议事日程，'
                           '建立健全网络安全管理制度，明确责任分工，确保各项措施落实到位。')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # 二级标题
    doc.add_heading('（一）强化组织领导', level=2)

    p = doc.add_paragraph('成立网络安全工作领导小组，由主要负责同志担任组长，'
                           '分管负责同志担任副组长，相关部门负责人为成员。')

    doc.add_heading('（二）完善制度建设', level=2)

    p = doc.add_paragraph('建立网络安全管理制度体系，包括网络安全策略、操作规程、'
                           '应急预案等，确保各项工作有章可循。')

    # 三级标题
    doc.add_heading('1. 制度制定要求', level=3)

    p = doc.add_paragraph('制度制定应符合国家相关法律法规和标准规范，结合本单位实际情况，'
                           '具有可操作性和针对性。')

    doc.add_heading('二、主要工作任务', level=1)

    p = doc.add_paragraph('（一）开展网络安全等级保护工作。按照《网络安全等级保护基本要求》，'
                           '对重要信息系统进行定级备案和安全建设整改。')

    p = doc.add_paragraph('（二）加强人员安全意识培训。定期组织网络安全知识培训，'
                           '提高全员网络安全意识和技能水平。')

    p = doc.add_paragraph('（三）建立安全监测预警机制。部署网络安全监测设备，'
                           '实时掌握网络安全态势，及时处置安全事件。')

    doc.add_heading('三、工作要求', level=1)

    p = doc.add_paragraph('各部门要认真落实本通知要求，于本月底前将工作方案报送网络安全工作领导小组办公室。'
                           '工作中遇到问题，请及时与办公室联系。')

    p = doc.add_paragraph('')
    p = doc.add_paragraph('                                          某某单位办公室')
    for run in p.runs:
        run.font.name = 'Arial'

    p = doc.add_paragraph('                                          2026年4月2日')

    doc.save(output_path)
    print(f'✓ 公文测试文档已创建：{output_path}')


def main():
    parser = argparse.ArgumentParser(description='生成格式测试用 Word 文档')
    parser.add_argument('--type', required=True, choices=['academic', 'gongwen'],
                        help='文档类型')
    parser.add_argument('--output', required=True, help='输出文件路径')
    args = parser.parse_args()

    if args.type == 'academic':
        create_academic_test_doc(args.output)
    else:
        create_gongwen_test_doc(args.output)


if __name__ == '__main__':
    main()
